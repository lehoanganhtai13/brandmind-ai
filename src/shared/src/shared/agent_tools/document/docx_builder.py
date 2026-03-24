"""DOCX builder for brand strategy documents using python-docx.

Generates editable Word documents with branded styles, TOC field,
section rendering, tables, and image embedding.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from loguru import logger

from .templates.brand_strategy import BrandStrategyTemplate


def _hex_to_rgb_color(hex_color: str) -> RGBColor:
    """Convert hex color string to RGBColor."""
    hex_color = hex_color.lstrip("#")
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )


class BrandStrategyDOCXBuilder:
    """Builds professional branded DOCX documents.

    Uses python-docx with Calibri font and brand colors applied
    to heading styles. Supports TOC fields, tables, and images.
    """

    def __init__(self, template: BrandStrategyTemplate | None = None) -> None:
        self.template = template or BrandStrategyTemplate()
        self.primary_color = _hex_to_rgb_color(self.template.brand_colors[0])

    def build(
        self,
        content: dict[str, Any],
        output_path: str,
        images: dict[str, str] | None = None,
    ) -> str:
        """Build DOCX document from content dict.

        Args:
            content: Dict with section content keyed by content_key.
            output_path: File path for output DOCX.
            images: Optional dict mapping section_id to image file path.

        Returns:
            Path to generated DOCX file.
        """
        doc = Document()
        self._apply_base_styles(doc)

        self._render_cover(doc)
        self._insert_toc_field(doc)

        for section in self.template.sections:
            if section.id == "cover":
                continue

            section_content = self._resolve_content(content, section.content_key)
            if section_content is None and not section.required:
                continue

            if section.page_break_before:
                doc.add_page_break()

            self._render_heading(doc, section.title, level=1)
            if section.subtitle:
                p = doc.add_paragraph(section.subtitle)
                p.style = doc.styles["Subtitle"]

            self._render_content(doc, section_content)

            if images and section.id in images:
                self._embed_image(doc, images[section.id])

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f"DOCX generated: {output_path}")
        return output_path

    def _apply_base_styles(self, doc: Any) -> None:
        """Set Calibri font and brand color headings."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        for level in range(1, 4):
            heading_style = doc.styles[f"Heading {level}"]
            heading_style.font.color.rgb = self.primary_color
            heading_style.font.name = "Calibri"

    def _render_cover(self, doc: Any) -> None:
        """Create cover page with brand name and title."""
        for _ in range(6):
            doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.template.brand_name)
        run.bold = True
        run.font.size = Pt(36)
        run.font.color.rgb = self.primary_color

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Brand Strategy Document")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(128, 128, 128)

    def _insert_toc_field(self, doc: Any) -> None:
        """Insert Word TOC field that auto-updates on open."""
        doc.add_page_break()
        self._render_heading(doc, "Table of Contents", level=1)

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)

        run2 = paragraph.add_run()
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2._r.append(instr_text)

        run3 = paragraph.add_run()
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fld_char_end)

    def _render_heading(self, doc: Any, text: str, level: int = 1) -> None:
        """Add heading with brand color."""
        doc.add_heading(text, level=level)

    def _render_content(self, doc: Any, content: Any) -> None:
        """Dispatch content rendering based on type."""
        if content is None:
            p = doc.add_paragraph("(No content available)")
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        elif isinstance(content, str):
            doc.add_paragraph(content)
        elif isinstance(content, dict):
            self._render_dict(doc, content)
        elif isinstance(content, list):
            if content and isinstance(content[0], dict):
                self._render_table(doc, content)
            else:
                for item in content:
                    doc.add_paragraph(str(item), style="List Bullet")

    def _render_dict(self, doc: Any, data: dict[str, Any], level: int = 2) -> None:
        """Render dict as hierarchical headings."""
        for key, value in data.items():
            label = key.replace("_", " ").title()
            heading_level = min(level, 4)
            doc.add_heading(label, level=heading_level)

            if isinstance(value, str):
                doc.add_paragraph(value)
            elif isinstance(value, dict):
                self._render_dict(doc, value, level + 1)
            elif isinstance(value, list):
                if value and isinstance(value[0], dict):
                    self._render_table(doc, value)
                else:
                    for item in value:
                        doc.add_paragraph(str(item), style="List Bullet")

    def _render_table(self, doc: Any, data: list[dict[str, Any]]) -> None:
        """Create formatted table from list of dicts."""
        if not data:
            return

        headers = list(data[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header.replace("_", " ").title()
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for row_data in data:
            row = table.add_row()
            for i, header in enumerate(headers):
                row.cells[i].text = str(row_data.get(header, ""))

    def _embed_image(self, doc: Any, image_path: str) -> None:
        """Embed centered image with optional caption."""
        path = Path(image_path)
        if not path.exists():
            logger.warning(f"Image not found: {image_path}")
            return

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(5.5))

    def _resolve_content(self, content: dict[str, Any], content_key: str) -> Any:
        """Resolve dotted content key path."""
        parts = content_key.split(".")
        current = content
        for part in parts:
            if isinstance(current, dict) and part in current:
                current = current[part]
            else:
                return None
        return current
